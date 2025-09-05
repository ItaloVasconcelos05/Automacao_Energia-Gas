from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
import os
import pandas as pd
from utils import (
    adicionar_paragrafo,
    adicionar_titulo_secao,
    aplicar_borda_paragrafo,
    adicionar_legenda_formatada,
    processar_imagem_para_relatorio,
)


def gerar_secao_nao_conformidades_constatadas(
    doc, row, nao_conformidades_df, fotos_dir, observacoes_df
):
    """
    Gera a seção '3. NÃO CONFORMIDADES CONSTATADAS', agora agrupada por terminal fiscalizado,
    com títulos de terminal e NCs em preto e negrito.
    """

    adicionar_titulo_secao(doc, "3. NÃO CONFORMIDADES CONSTATADAS")

    adicionar_paragrafo(
        doc,
        "A seguir, apresentam-se as não conformidades registradas nos diversos terminais fiscalizados:",
    )

    id_fisc = row["ID_FISC"]

    nc_fisc = nao_conformidades_df[
        nao_conformidades_df["ID_FISC_original"] == id_fisc
    ]

    if "Terminal" not in nc_fisc.columns:
        adicionar_paragrafo(
            doc, "⚠️ Coluna 'Terminal' não encontrada na planilha de não conformidades."
        )
        return

    # Agrupar por terminal
    for terminal, grupo_terminal in nc_fisc.groupby("Terminal"):
        # Título do terminal (nível 2) — usar formatação igual à adicionar_titulo_secao e tudo em maiúsculo
        par_terminal = doc.add_paragraph()
        par_terminal.add_run(f"{terminal.upper()}").bold = True

        # Agrupar por número da não conformidade dentro de cada terminal
        for nc_id, grupo_nc in grupo_terminal.groupby("Nº"):
            descricao = grupo_nc["Não Conformidade"].iloc[0]

            # Título da NC (nível 3) — também em negrito, preto e agora sublinhado
            par_nc = doc.add_paragraph()
            run_nc = par_nc.add_run(f"{nc_id} - {descricao}")
            run_nc.bold = True
            run_nc.underline = True
            run_nc.font.size = Pt(10)
            run_nc.font.color.rgb = RGBColor(0, 0, 0)

            for _, linha in grupo_nc.iterrows():
                # Supondo que o separador seja ';'
                nomes_fotos = (
                    [f.strip() for f in str(linha["Foto"]).split(";") if f.strip()]
                    if pd.notna(linha["Foto"])
                    else []
                )
                legendas = (
                    [l.strip() for l in str(linha["Legenda da Foto"]).split(";")]
                    if pd.notna(linha["Legenda da Foto"])
                    else []
                )

                for idx, nome_foto in enumerate(nomes_fotos):
                    foto_path = os.path.join(fotos_dir, nome_foto)
                    legenda = legendas[idx] if idx < len(legendas) else ""
                    if os.path.exists(foto_path):
                        buffer = processar_imagem_para_relatorio(foto_path)
                        doc.add_picture(buffer, width=Inches(3))
                        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        aplicar_borda_paragrafo(doc.paragraphs[-1])
                        adicionar_legenda_formatada(doc, legenda)
                    else:
                        if legenda:
                            adicionar_paragrafo(doc, legenda)

        # 🔽 OBSERVAÇÕES IMPORTANTES PARA O TERMINAL
        obs_terminais = observacoes_df[
            (observacoes_df["ID_FISC"] == id_fisc)
            & (observacoes_df["Terminal"] == terminal)
        ]

        if not obs_terminais.empty:
            adicionar_titulo_secao(doc, "OBSERVAÇÕES IMPORTANTES")
            # Deixar o último parágrafo (título) sublinhado, maiúsculo e tamanho igual ao da NC
            par_titulo_obs = doc.paragraphs[-1]
            run_titulo_obs = par_titulo_obs.runs[0]
            run_titulo_obs.text = run_titulo_obs.text.upper()
            run_titulo_obs.underline = True
            run_titulo_obs.font.size = Pt(10)

            for _, obs in obs_terminais.iterrows():
                texto_obs = (
                    str(obs["Observações"]).strip()
                    if pd.notna(obs["Observações"])
                    else ""
                )
                # Supondo que o separador seja ';'
                nomes_fotos_obs = (
                    [f.strip() for f in str(obs["Foto"]).split(";") if f.strip()]
                    if pd.notna(obs["Foto"])
                    else []
                )
                legendas_obs = (
                    [l.strip() for l in str(obs["Legenda da Foto"]).split(";")]
                    if pd.notna(obs["Legenda da Foto"])
                    else []
                )

                if texto_obs:
                    adicionar_paragrafo(doc, texto_obs)

                for idx, foto_obs in enumerate(nomes_fotos_obs):
                    foto_path = os.path.join(fotos_dir, foto_obs)
                    legenda_obs = legendas_obs[idx] if idx < len(legendas_obs) else ""
                    if os.path.exists(foto_path):
                        buffer = processar_imagem_para_relatorio(foto_path)
                        doc.add_picture(buffer, width=Inches(3))
                        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        aplicar_borda_paragrafo(doc.paragraphs[-1])
                        if legenda_obs:
                            adicionar_legenda_formatada(doc, legenda_obs)
                    elif legenda_obs:
                        adicionar_legenda_formatada(doc, legenda_obs)

    ## doc.add_section(WD_SECTION.NEW_PAGE)
