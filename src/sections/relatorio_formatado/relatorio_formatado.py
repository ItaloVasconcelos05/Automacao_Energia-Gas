# relatorio_formatado.py

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
import pandas as pd
from utils import (
    adicionar_titulo_secao,
    adicionar_paragrafo_justificado,
    ajustar_largura_colunas,
    arquivo_em_uso,
)


def gerar_secao_nao_conformidades_constatadas(
    doc, row, nao_conformidades_df, fotos_dir, observacoes_df
):
    adicionar_titulo_secao(doc, "3. NAO CONFORMIDADES CONSTATADAS")

    id_fisc = row["ID da Fiscalização"]
    nc_fisc = nao_conformidades_df[
        nao_conformidades_df["ID da Fiscalização"] == id_fisc
    ]

    if "Terminal" not in nc_fisc.columns:
        doc.add_paragraph("⚠️ Coluna 'Terminal' ausente na planilha.")
        return

    for terminal, grupo_terminal in nc_fisc.groupby("Terminal"):
        par_terminal = doc.add_paragraph()
        run_terminal = par_terminal.add_run(terminal.upper())
        aplicar_estilo_texto(run_terminal, tamanho=14, negrito=True)

        for nc_id, grupo_nc in grupo_terminal.groupby("Nº"):
            descricao = grupo_nc["Não Conformidade"].iloc[0]
            par_nc = doc.add_paragraph()
            run_nc = par_nc.add_run(f"{nc_id} - {descricao}")
            aplicar_estilo_texto(run_nc, tamanho=12, negrito=True)

            imagens = []
            for _, linha in grupo_nc.iterrows():
                nome_foto = str(linha.get("Foto", "")).strip()
                legenda = str(linha.get("Legenda da Foto", "")).strip()
                foto_path = os.path.join(fotos_dir, nome_foto)

                if nome_foto and os.path.exists(foto_path):
                    imagens.append((foto_path, legenda))
                elif legenda:
                    adicionar_paragrafo_justificado(doc, legenda)

            for i, (foto_path, legenda) in enumerate(imagens):
                doc.add_picture(foto_path, width=Inches(2.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                adicionar_legenda_formatada(doc, legenda)

        # Observações importantes
        obs_terminais = observacoes_df[
            (observacoes_df["ID da Fiscalização"] == id_fisc)
            & (observacoes_df["Terminal"] == terminal)
        ]

        if not obs_terminais.empty:
            adicionar_titulo_secao(doc, "Observações Importantes")

            for _, obs in obs_terminais.iterrows():
                texto_obs = (
                    str(obs["Observações"]).strip()
                    if pd.notna(obs["Observações"])
                    else ""
                )
                foto_obs = str(obs["Foto"]).strip() if pd.notna(obs["Foto"]) else ""
                legenda_obs = (
                    str(obs["Legenda da Foto"]).strip()
                    if pd.notna(obs["Legenda da Foto"])
                    else ""
                )

                if texto_obs:
                    adicionar_paragrafo_justificado(doc, texto_obs)

                if foto_obs and os.path.exists(os.path.join(fotos_dir, foto_obs)):
                    doc.add_picture(
                        os.path.join(fotos_dir, foto_obs), width=Inches(2.5)
                    )
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if legenda_obs:
                        adicionar_legenda_formatada(doc, legenda_obs)
                elif legenda_obs:
                    adicionar_legenda_formatada(doc, legenda_obs)

    doc.add_section(WD_SECTION.NEW_PAGE)
