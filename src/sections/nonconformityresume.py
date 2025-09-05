from utils import adicionar_titulo_secao
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION


def gerar_secao_resumo_nao_conformidades(doc, row, nao_conformidades_df):
    """
    Gera a seção '4. RESUMO DAS NÃO CONFORMIDADES IDENTIFICADAS' no formato visual do relatório oficial,
    com agrupamento por terminal e número.

    Parâmetros:
    - doc: objeto Document.
    - row: linha da fiscalização (pandas.Series).
    - nao_conformidades_df: DataFrame com as colunas: 'ID da Fiscalização', 'Terminal', 'Nº', 'Não Conformidade'.
    """

    espaco1 = doc.add_paragraph()
    espaco1.paragraph_format.space_after = Pt(12)

    adicionar_titulo_secao(doc, "4. RESUMO DAS NÃO CONFORMIDADES IDENTIFICADAS")

    id_fisc = row["ID_FISC"]

    nc_fisc = nao_conformidades_df[
        nao_conformidades_df["ID_FISC_original"] == id_fisc
    ]

    if nc_fisc.empty:
        doc.add_paragraph("Nenhuma não conformidade registrada.")
        return

    if "Terminal" not in nc_fisc.columns or "Nº" not in nc_fisc.columns:
        doc.add_paragraph("⚠️ Colunas obrigatórias não encontradas na planilha.")
        return

    # Criar a tabela com 2 colunas
    tabela = doc.add_table(rows=1, cols=2)
    tabela.style = "Table Grid"
    tabela.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Cabeçalhos
    cabecalho = tabela.rows[0].cells
    cabecalho[0].text = "TERMINAL"
    cabecalho[1].text = "NÃO CONFORMIDADE"

    # Aplicar estilo ao cabeçalho
    for cell in cabecalho:
        for par in cell.paragraphs:
            run = par.runs[0]
            run.bold = True
            run.font.size = Pt(11)

    # Agrupar por Terminal e listar NCs numeradas
    for terminal, grupo in nc_fisc.groupby("Terminal"):
        grupo = grupo.sort_values(by="Nº")  # ordena as NCs

        for idx, (_, linha) in enumerate(grupo.iterrows()):
            row_cells = tabela.add_row().cells

            # Primeira linha do grupo: escreve terminal
            if idx == 0:
                row_cells[0].text = f"{terminal}"
            else:
                row_cells[0].text = ""  # manter célula vazia nas linhas seguintes

            # Coluna da NC formatada com número + descrição
            numero_nc = str(linha["Nº"]).strip()
            titulo = f"Não Conformidade {numero_nc}"
            descricao = linha["Não Conformidade"].strip()
            run = row_cells[1].paragraphs[0].add_run(f"{titulo} – {descricao}")
            run.bold = True
            run.font.size = Pt(11)

    espaco = doc.add_paragraph()
    espaco.paragraph_format.space_after = Pt(24)

    ## doc.add_section(WD_SECTION.NEW_PAGE)
