from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
import os
import pandas as pd
from utils import (
    adicionar_paragrafo_justificado,
    adicionar_titulo_secao,
    adicionar_texto_centralizado,
    aplicar_borda_paragrafo,  # importar função
    adicionar_legenda_formatada,  # importar função
)


def gerar_secao_nao_conformidades_constatadas(
    doc, row, nao_conformidades_df, fotos_dir, observacoes_df
):
    """
    Gera a seção '3. NÃO CONFORMIDADES CONSTATADAS', agora agrupada por terminal fiscalizado,
    com títulos de terminal e NCs em preto e negrito.
    """

    adicionar_titulo_secao(doc, "3. NÃO CONFORMIDADES CONSTATADAS")

    adicionar_paragrafo_justificado(
        doc,
        "A seguir, apresentam-se as não conformidades registradas nos diversos terminais fiscalizados:",
    )

    id_fisc = row["ID da Fiscalização"]

    nc_fisc = nao_conformidades_df[
        nao_conformidades_df["ID da Fiscalização"] == id_fisc
    ]

    if "Terminal" not in nc_fisc.columns:
        adicionar_paragrafo_justificado(
            doc, "⚠️ Coluna 'Terminal' não encontrada na planilha de não conformidades."
        )
        return

    # Agrupar por terminal
    for terminal, grupo_terminal in nc_fisc.groupby("Terminal"):
        # Título do terminal (nível 2) — usar formatação igual à adicionar_titulo_secao e tudo em maiúsculo
        par_terminal = doc.add_paragraph()
        par_terminal.add_run(f"{terminal.upper()}").bold = True

        # Agrupar por número da não conformidade dentro de cada terminal
        for nc_id, grupo_nc in grupo_terminal.groupby("Nº"):
            descricao = grupo_nc["Não Conformidade"].iloc[0]

            # Título da NC (nível 3) — também em negrito, preto e agora sublinhado
            par_nc = doc.add_paragraph()
            run_nc = par_nc.add_run(f"{nc_id} - {descricao}")
            run_nc.bold = True
            run_nc.underline = True
            run_nc.font.size = Pt(10)
            run_nc.font.color.rgb = RGBColor(0, 0, 0)

            for _, linha in grupo_nc.iterrows():
                nome_foto = str(linha["Foto"]) if pd.notna(linha["Foto"]) else ""
                legenda = (
                    str(linha["Legenda da Foto"])
                    if pd.notna(linha["Legenda da Foto"])
                    else ""
                )
                foto_path = os.path.join(fotos_dir, nome_foto) if nome_foto else ""

                if nome_foto and os.path.exists(foto_path):
                    doc.add_picture(foto_path, width=Inches(3))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    aplicar_borda_paragrafo(doc.paragraphs[-1])
                    adicionar_legenda_formatada(doc, legenda)
                else:
                    if legenda:
                        adicionar_paragrafo_justificado(doc, legenda)

        # 🔽 OBSERVAÇÕES IMPORTANTES PARA O TERMINAL
        obs_terminais = observacoes_df[
            (observacoes_df["ID da Fiscalização"] == id_fisc)
            & (observacoes_df["Terminal"] == terminal)
        ]

        if not obs_terminais.empty:
            adicionar_titulo_secao(doc, "OBSERVAÇÕES IMPORTANTES")
            # Deixar o último parágrafo (título) sublinhado, maiúsculo e tamanho igual ao da NC
            par_titulo_obs = doc.paragraphs[-1]
            run_titulo_obs = par_titulo_obs.runs[0]
            run_titulo_obs.text = run_titulo_obs.text.upper()
            run_titulo_obs.underline = True
            run_titulo_obs.font.size = Pt(10)

            for _, obs in obs_terminais.iterrows():
                texto_obs = (
                    str(obs["Observações"]).strip()
                    if pd.notna(obs["Observações"])
                    else ""
                )
                foto_obs = str(obs["Foto"]).strip() if pd.notna(obs["Foto"]) else ""
                legenda_obs = (
                    str(obs["Legenda da Foto"]).strip()
                    if pd.notna(obs["Legenda da Foto"])
                    else ""
                )
                foto_path = os.path.join(fotos_dir, foto_obs) if foto_obs else ""

                if texto_obs:
                    adicionar_paragrafo_justificado(doc, texto_obs)

                if foto_obs and os.path.exists(foto_path):
                    doc.add_picture(foto_path, width=Inches(3))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    aplicar_borda_paragrafo(doc.paragraphs[-1])
                    if legenda_obs:
                        adicionar_legenda_formatada(doc, legenda_obs)
                elif legenda_obs:
                    adicionar_legenda_formatada(doc, legenda_obs)

    ## doc.add_section(WD_SECTION.NEW_PAGE)
