# Adiciona importação necessária para WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_ALIGN_PARAGRAPH
# Função para adicionar parágrafo com marcador, título em negrito e texto normal
def adicionar_paragrafo_bullet_bold(doc, titulo, texto, tamanho_fonte=12, font_name='Calibri', alinhamento=WD_ALIGN_PARAGRAPH.JUSTIFY, espaco_antes=0, espaco_depois=0):
    paragraph = doc.add_paragraph(style='List Bullet')
    run_titulo = paragraph.add_run(titulo)
    run_titulo.bold = True
    run_titulo.font.size = Pt(tamanho_fonte)
    run_titulo.font.name = font_name
    run_texto = paragraph.add_run(f" {texto}")
    run_texto.bold = False
    run_texto.font.size = Pt(tamanho_fonte)
    run_texto.font.name = font_name
    paragraph.alignment = alinhamento
    paragraph.paragraph_format.space_before = Pt(espaco_antes)
    paragraph.paragraph_format.space_after = Pt(espaco_depois)
    return paragraph
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from openpyxl import load_workbook
import os
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image
import io
from typing import Literal

ESTILOS_COMUNS = Literal['Normal', 'Heading 1', 'Heading 2', 'Heading 3', 'List Bullet', 'List Number', 'Title', 'Subtitle', 'Quote', 'Caption', 'Body Text', 'Footnote Text', 'Endnote Text', 'Table Grid', 'Hyperlink']

# Funções auxiliares de formatação:
def adicionar_paragrafo(doc, texto, 
                        tamanho_fonte=12, 
                        estilo: ESTILOS_COMUNS = 'Normal',
                        bold = False, 
                        font_name='Calibri', 
                        alinhamento=WD_ALIGN_PARAGRAPH.LEFT, 
                        espaco_antes=0, 
                        espaco_depois=0,
                        cor=None):
    """
    Adiciona um parágrafo com texto, fonte, tamanho e alinhamento personalizáveis.
    
    Args:
        doc (Document): O objeto do documento python-docx.
        texto (str): O texto a ser adicionado.  
        tamanho_fonte (int): Tamanho da fonte do texto (opcional, padrão 12).
        estilo (ESTILOS_COMUNS): Estilo do parágrafo (opcional, padrão None).
        bold (bool): Se o texto deve ser negrito (opcional, padrão False).
        font_name (str): Nome da fonte do texto (opcional, padrão 'Calibri').
        alinhamento (WD_ALIGN_PARAGRAPH): O alinhamento do parágrafo (LEFT, CENTER, RIGHT, JUSTIFY). Padrão é LEFT.
        espaco_antes (int): Espaçamento acima do parágrafo em pontos (Pt). Padrão é 0.
        espaco_depois (int): Espaçamento abaixo do parágrafo em pontos (Pt). Padrão é 0.
        cor (tuple of int, optional): A cor do texto em formato RGB (ex: (255, 0, 0) para vermelho). Padrão é None.
    """
    paragraph = doc.add_paragraph(style=estilo)
    run = paragraph.add_run(texto)
    if cor is not None and len(cor) == 3:
        run.font.color.rgb = RGBColor(cor[0], cor[1], cor[2])
    run.font.name = font_name
    run.font.size = Pt(tamanho_fonte)
    run.bold = bold
    paragraph.alignment = alinhamento
    paragraph.paragraph_format.space_before = Pt(espaco_antes)
    paragraph.paragraph_format.space_after = Pt(espaco_depois)
    
    return paragraph

def adicionar_tabela(doc, dados, cabecalho=None, largura_colunas=None, estilo_tabela='Table Grid', 
                     alinhamento_cabecalho=WD_ALIGN_PARAGRAPH.CENTER, 
                     alinhamento_colunas=None):
    """
    Adiciona uma tabela padronizada a um documento Word.

    Args:
        doc (Document): O objeto do documento python-docx.
        dados (list of list/tuple): Os dados da tabela. Cada lista/tupla representa uma linha.
        cabecalho (list of str, optional): O cabeçalho da tabela. Se None, o primeiro item de 'dados' será usado como cabeçalho.
        largura_colunas (list of int/float, optional): A largura de cada coluna em polegadas.
        estilo_tabela (str): O estilo da tabela a ser aplicado. Padrão é 'Table Grid'.
        alinhamento_cabecalho (WD_ALIGN_PARAGRAPH): O alinhamento horizontal do texto no cabeçalho.
        alinhamento_colunas (list of WD_ALIGN_PARAGRAPH, optional): Uma lista com o alinhamento horizontal para cada coluna.
    """
    if not dados:
        return

    # Se não houver cabeçalho, a primeira linha dos dados é usada
    if cabecalho is None:
        cabecalho = dados[0]
        dados_corpo = dados[1:]
    else:
        dados_corpo = dados
    
    num_cols = len(cabecalho)
    tabela = doc.add_table(rows=1, cols=num_cols)
    tabela.style = estilo_tabela

    # Define a largura das colunas se o argumento for fornecido
    if largura_colunas and len(largura_colunas) == num_cols:
        for i, width in enumerate(largura_colunas):
            tabela.columns[i].width = Inches(width)

    # Preenche o cabeçalho
    hdr_cells = tabela.rows[0].cells
    for i, cell_text in enumerate(cabecalho):
        paragraph = hdr_cells[i].paragraphs[0]
        hdr_cells[i].text = cell_text
        hdr_cells[i].paragraphs[0].alignment = alinhamento_cabecalho
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for run in paragraph.runs:
            run.bold = True
                
    # Preenche o corpo da tabela
    for linha in dados_corpo:
        cells = tabela.add_row().cells
        for i, cell_text in enumerate(linha):
            cells[i].text = str(cell_text) # Garante que o texto seja uma string
            
            # Alinha o texto das células do corpo
            if alinhamento_colunas and len(alinhamento_colunas) == num_cols:
                cells[i].paragraphs[0].alignment = alinhamento_colunas[i]
            else:
                cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT # Alinhamento padrão
            
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def adicionar_imagem_no_cabecalho(doc, imagem_path, largura=Inches(1), alinhamento=WD_PARAGRAPH_ALIGNMENT.LEFT):
    """
    Adiciona uma imagem APENAS ao cabeçalho da primeira página da primeira seção do documento.

    Args:
        doc (Document): O objeto do documento python-docx.
        imagem_path (str): O caminho para o arquivo da imagem.
        largura (Inches): A largura da imagem (opcional, padrão Inches(1)).
        alinhamento (WD_PARAGRAPH_ALIGNMENT): O alinhamento da imagem no cabeçalho (opcional, padrão LEFT).
    """
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    first_page_header = section.first_page_header
    if not first_page_header.paragraphs:
        header_paragraph = first_page_header.add_paragraph()
    else:
        header_paragraph = first_page_header.paragraphs[0]
    header_paragraph.clear()
    run = header_paragraph.add_run()
    run.add_picture(imagem_path, width=largura)
    header_paragraph.alignment = alinhamento
    header_paragraph.paragraph_format.space_before = Pt(0)
    header_paragraph.paragraph_format.space_after = Pt(0)

def adicionar_imagem_paragrafo(paragrafo, imagem_path, largura=Inches(1)):
    """
    Insere uma imagem centralizada em um parágrafo existente.

    Args:
        doc (Document): O objeto do documento python-docx.
        paragrafo (Paragraph): O parágrafo onde a imagem será inserida.
        imagem_path (str): O caminho para o arquivo da imagem.
        largura (Inches): A largura desejada para a imagem.
    """
    run = paragrafo.add_run()
    run.add_picture(imagem_path, width=largura)
    paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

def adicionar_titulo_secao(doc, texto):
    """Adiciona um título de seção formatado."""

    secao = doc.add_paragraph()
    secao.add_run(texto).bold = True


# Função para ajustar a largura das colunas
def ajustar_largura_colunas(caminho_planilha):
    wb = load_workbook(caminho_planilha)
    ws = wb.active

    for coluna in ws.columns:
        max_length = 0
        coluna_letra = coluna[0].column_letter

        for celula in coluna:
            try:
                if celula.value:
                    max_length = max(max_length, len(str(celula.value)))
            except:
                pass

        # Define largura da coluna com margem extra
        ajuste = max_length + 2
        ws.column_dimensions[coluna_letra].width = ajuste

    wb.save(caminho_planilha)


# Função para verificar se arquivo está em uso
def arquivo_em_uso(caminho):
    try:
        os.rename(caminho, caminho)
        return False
    except PermissionError:
        return True


def aplicar_estilo_texto(
    run, tamanho=12, negrito=False, fonte="Arial", cor_rgb=(0, 0, 0)
):
    run.font.name = fonte
    run._element.rPr.rFonts.set(qn("w:eastAsia"), fonte)
    run.font.size = Pt(tamanho)
    run.bold = negrito
    run.font.color.rgb = RGBColor(*cor_rgb)


def aplicar_borda_paragrafo(paragraph):
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    for border_name in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "2")
        border.set(qn("w:color"), "000000")
        borders.append(border)
    pPr.append(borders)


def adicionar_legenda_formatada(doc, texto):
    par = doc.add_paragraph()
    run = par.add_run(texto)
    aplicar_estilo_texto(run, tamanho=10, fonte="Arial", cor_rgb=(90, 90, 90))
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    aplicar_borda_paragrafo(par)


def processar_imagem_para_relatorio(caminho_imagem, largura_max=1024, qualidade=80):
    # Abre a imagem
    img = Image.open(caminho_imagem)
    # Converte para RGB se necessário (evita problemas com PNG/transparência)
    if img.mode != "RGB":
        img = img.convert("RGB")
    # Redimensiona mantendo proporção
    if img.width > largura_max:
        proporcao = largura_max / float(img.width)
        altura_nova = int(float(img.height) * proporcao)
        img = img.resize((largura_max, altura_nova), Image.LANCZOS)
    # Salva em memória, sem metadados, com compressão JPEG
    buffer = io.BytesIO()
    img.save(buffer, format="JPEG", quality=qualidade, optimize=True)
    buffer.seek(0)
    return buffer
