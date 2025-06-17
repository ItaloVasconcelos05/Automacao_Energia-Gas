import pandas as pd
from docx import Document
from docx.shared import Inches
import os
from docx2pdf import convert
from tqdm import tqdm
import sys
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from utils import (
    adicionar_paragrafo_justificado,
    adicionar_texto_centralizado,
    adicionar_titulo_secao,
    ajustar_largura_colunas,
    arquivo_em_uso,
)


def main():
    # === CONFIGURAÇÃO DE CAMINHOS BASE ===
    # Obtém o diretório atual (compatível com execução como .exe ou .py)

    if getattr(sys, "frozen", False):
        BASE_DIR = os.path.dirname(sys.executable)
    else:
        BASE_DIR = os.path.dirname(os.path.abspath(__file__))

    # Configurações
    FOTOS_DIR = os.path.join(BASE_DIR, "assets")
    RELATORIOS_DIR = os.path.join(BASE_DIR, "reports")
    CAMINHO_PLANILHA = os.path.join(BASE_DIR, "planilha_fiscalizacao.xlsx")
    COLUNA_STATUS = "Relatório Gerado"

    # Cria pasta de relatórios se não existir
    os.makedirs(RELATORIOS_DIR, exist_ok=True)

    # Criar pasta de fotos se não existir
    os.makedirs(FOTOS_DIR, exist_ok=True)

    # Verifica se a planilha está em uso
    if arquivo_em_uso(CAMINHO_PLANILHA):
        print(
            f"⚠️ ERRO: O arquivo '{CAMINHO_PLANILHA}' está aberto ou em uso. Feche-o e execute novamente."
        )
        exit(1)

    # Carrega planilha
    planilha = pd.read_excel(CAMINHO_PLANILHA)

    # Garante que a coluna de status existe e é booleana
    if COLUNA_STATUS not in planilha.columns:
        planilha[COLUNA_STATUS] = False

    planilha[COLUNA_STATUS] = planilha[COLUNA_STATUS].astype(bool)

    # Processa apenas fiscalizações sem relatório
    pendentes = planilha[~planilha[COLUNA_STATUS]]

    if pendentes.empty:
        print("✅ Todos os relatórios já foram gerados. Nada a fazer.")
        return

    # Itera sobre fiscalizações pendentes
    for idx_pendente in tqdm(
        pendentes.index, total=pendentes.shape[0], desc="Gerando relatórios"
    ):
        row = planilha.loc[idx_pendente]

        # --- INÍCIO DA LÓGICA DE GERAÇÃO DO RELATÓRIO ---
        # Substitua o conteúdo do seu loop 'for' por este bloco

        # Assumindo que seu loop original é:
        # for index, row in tqdm(pendentes.iterrows(), ...):

        doc = Document()

        # =================================================================
        # 1. SEÇÃO DA CAPA
        # =================================================================
        # Adicionar o logo da ARPE (substitua 'logo_arpe.png' pelo caminho do seu arquivo)
        # doc.add_picture('logo_arpe.png', width=Inches(2.0))
        # doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

        adicionar_texto_centralizado(
            doc, "\n\nRELATÓRIO DE FISCALIZAÇÃO CTR 02/2024", tamanho_fonte=12
        )
        adicionar_texto_centralizado(
            doc, "\nTerminais Rodoviários Intermunicipais Concedidos", tamanho_fonte=12
        )

        adicionar_texto_centralizado(
            doc,
            "\nRecife (TIP), Caruaru, Arcoverde, Garanhuns, Serra Talhada e Petrolina",
        )
        adicionar_texto_centralizado(
            doc, "Contrato de Concessão de Serviço Público Nº 1.041.080/08"
        )

        adicionar_texto_centralizado(
            doc, "\n\n\nCOORDENADORIA DE TRANSPORTES E RODOVIAS"
        )
        adicionar_texto_centralizado(
            doc, "Coordenadora: Maria Ângela Albuquerque de Freitas"
        )
        adicionar_texto_centralizado(
            doc, "Analista de Regulação: Enildo Manoel da Silva Junior"
        )
        adicionar_texto_centralizado(
            doc, "Assistente de Regulação e Fiscalização: Domingos Sávio Menezes"
        )

        adicionar_texto_centralizado(doc, "\n\n\n\nAbril de 2024")

        # Mudar para a próxima página (nova seção)
        doc.add_section(WD_SECTION.NEW_PAGE)

        # =================================================================
        # 2. SEÇÃO DE OBJETIVOS E LEGISLAÇÃO
        # =================================================================
        adicionar_texto_centralizado(doc, "RELATÓRIO DE FISCALIZAÇÃO")
        adicionar_texto_centralizado(doc, "CTR 02/2024")

        adicionar_titulo_secao(doc, "\nI - OBJETIVOS")
        adicionar_paragrafo_justificado(
            doc,
            "Verificar as condições operacionais, de conservação, de manutenção e de segurança dos terminais intermunicipais, bem como do cumprimento da legislação aplicada e da eficiência do serviço.",
        )
        adicionar_paragrafo_justificado(
            doc,
            "A Coordenadoria de Transportes e Rodovias da Arpe, realizou no período de 26 de fevereiro a 1º de março de 2024, fiscalização nos Terminais Rodoviários Intermunicipais concedidos à Empresa SOCICAM - Administração, Projetos e Representações Ltda (SOCICAM), nas cidades de Recife – TIP, Caruaru, Garanhuns, Arcoverde, Serra Talhada e Petrolina, conforme Contrato de Serviço Público Nº 1.041.080/08, firmado entre o Governo do Estado, representado pela Secretaria de Transportes (SETRA) e a SOCICAM.",
        )

        adicionar_titulo_secao(doc, "\nII - LEGISLAÇÃO APLICADA")

        # Lista de legislação com formatação
        p = doc.add_paragraph(style="List Bullet")
        run1 = p.add_run("Lei nº 12.524, de 30 de dezembro de 2003")
        run1.bold = True
        run2 = p.add_run(
            " – Altera e consolida as disposições da Lei nº 12.126, de 12 de dezembro de 2001, que cria a Agência de Regulação dos Serviços Públicos do Estado de Pernambuco – ARPE, regulamentada pelo "
        )
        run3 = p.add_run("Decreto nº 30.200, de 09 de fevereiro de 2007;")
        run3.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = doc.add_paragraph(style="List Bullet")
        run1 = p.add_run("Lei nº 13.254, de 21 de junho de 2007")
        run1.bold = True
        run2 = p.add_run(
            " - Estrutura o Sistema de Transporte Coletivo Intermunicipal de Passageiros do Estado de Pernambuco, autoriza a criação da Empresa Pernambucana de Transportes Intermunicipal – EPTI, e alterações, em especial a "
        )
        run3 = p.add_run("Lei Estadual nº 15.200, de 17 de dezembro de 2013;")
        run3.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = doc.add_paragraph(style="List Bullet")
        run1 = p.add_run("Resolução ARPE nº 46, de 07 de abril de 2008")
        run1.bold = True
        run2 = p.add_run(
            " (Antiga 006/2008) - Aprova o Regulamento dos Terminais Rodoviários do Estado de Pernambuco, alterada parcialmente pela "
        )
        run3 = p.add_run("Resolução ARPE nº 53, de 26 de janeiro de 2009")
        run3.bold = True
        run4 = p.add_run(" (Antiga 003/2009);")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = doc.add_paragraph(style="List Bullet")
        run1 = p.add_run(
            "Contrato de Concessão de Serviço Público Nº 1.041.080/08, de 19 de setembro de 2008 e Termos Aditivos"
        )
        run1.bold = True
        run2 = p.add_run(
            " – Concessão de serviço público para operação, manutenção e administração de terminais rodoviários no Estado de Pernambuco, com execução de obras de reforma e construção, incluindo, ainda, a cessão de uso de espaços para a exploração comercial através de locação e publicidade, que entre si fazem, de um lado o Estado de Pernambuco, através da SETRA – Secretaria de Transportes – e de outro lado, a SOCICAM – Administração, projetos e representações;"
        )
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Mudar para a próxima página
        doc.add_section(WD_SECTION.NEW_PAGE)

        # =================================================================
        # 3. SEÇÃO DE METODOLOGIA E CONSTATAÇÕES
        # =================================================================
        adicionar_titulo_secao(doc, "III - METODOLOGIA")
        adicionar_paragrafo_justificado(
            doc,
            "Este relatório descreve as observações realizadas pela equipe da Coordenadoria de Transportes e Rodovias da Arpe, do ponto de vista técnico-operacional, registrando os aspectos mais relevantes.",
        )
        adicionar_paragrafo_justificado(
            doc,
            "Os procedimentos utilizados foram a verificação das condições dos serviços prestados em cada terminal rodoviário concedido à Empresa SOCICAM.",
        )
        adicionar_paragrafo_justificado(
            doc,
            "A ação fiscalizadora abrangeu toda a área dos terminais, verificação e análise de irregularidades e não conformidades, tomando por base o Contrato de Concessão de Serviço Público nº 1.041.080/2008.",
        )
        adicionar_paragrafo_justificado(
            doc,
            "Foram vistoriadas as condições de higiene das áreas de embarque e desembarque, os sanitários, as condições do pavimento das vias de circulação interna, a infraestrutura oferecida, os locais de estocagem de veículos, a segurança e o atendimento ao usuário.",
        )

        adicionar_titulo_secao(doc, "\nIV - PERÍODO DA FISCALIZAÇÃO E EQUIPE TÉCNICA")
        adicionar_paragrafo_justificado(
            doc,
            "A Fiscalização nos terminais foi realizada no período de 26 de fevereiro a 1º de março do ano em curso, pela equipe técnica formada pelo Analista de Regulação: Enildo Manoel da Silva Júnior e o Assistente de Regulação: Domingos Sávio Menezes, sob supervisão da Coordenadora de Transportes e Rodovias: Maria Ângela A. de Freitas.",
        )

        adicionar_titulo_secao(doc, "\nV - CONSTATAÇÕES")
        adicionar_paragrafo_justificado(
            doc,
            "Foram vistoriados seis terminais intermunicipais concedidos nas cidades de Recife (TIP), Caruaru, Garanhuns, Arcoverde, Serra Talhada e Petrolina.",
        )
        adicionar_paragrafo_justificado(
            doc,
            "As não conformidades encontradas estão relacionadas a seguir e também as principais ações nos terminais.",
        )

        # =================================================================
        # # # # # # # # # SEU CÓDIGO DINÂMICO ENTRA AQUI # # # # # # # # #
        # =================================================================
        # O seu código original que itera sobre as não conformidades e fotos
        # se encaixa perfeitamente aqui. Ele irá adicionar os problemas específicos
        # do `row` atual do seu loop `for`.

        # Exemplo de como seu código se encaixaria:
        doc.add_heading(f"{row['ID da Fiscalização']} - {row['Local']}", level=1)

        # Adicionando a descrição da não conformidade
        # Ex: 1.1- Verifica-se colunas da plataforma...
        doc.add_paragraph(f"{row['Não conformidade']}")

        # Adicionando as fotos
        fotos = str(row["Fotos"]).split(";") if pd.notnull(row["Fotos"]) else []
        for foto_legenda in fotos:
            # Supondo que a legenda esteja junto ao nome do arquivo, ex: "foto01.jpg|Legenda da foto"
            try:
                nome_foto, legenda = foto_legenda.split("|")
                foto_path = os.path.join(FOTOS_DIR, nome_foto.strip())
                if os.path.exists(foto_path):
                    doc.add_picture(
                        foto_path, width=Inches(3)
                    )  # Ajuste a largura conforme necessário
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    adicionar_texto_centralizado(doc, legenda.strip())
                else:
                    doc.add_paragraph(f"Foto não encontrada: {nome_foto.strip()}")
            except ValueError:
                # Lidar com casos onde não há legenda
                foto_path = os.path.join(FOTOS_DIR, foto_legenda.strip())
                if os.path.exists(foto_path):
                    doc.add_picture(foto_path, width=Inches(3))
                else:
                    doc.add_paragraph(f"Foto não encontrada: {foto_legenda.strip()}")

        # =================================================================
        # FIM DA SEÇÃO DE CÓDIGO DINÂMICO
        # =================================================================

        # Mudar para a próxima página
        doc.add_section(WD_SECTION.NEW_PAGE)

        # =================================================================
        # 4. SEÇÃO DE CONCLUSÕES E RECOMENDAÇÕES
        # =================================================================
        adicionar_titulo_secao(doc, "VII - CONCLUSÕES E RECOMENDAÇÕES")
        adicionar_paragrafo_justificado(
            doc,
            "Diante das constatações apontadas no presente Relatório, solicitamos um plano de ação para regularização das não conformidades com prazo para resolução, a fim de atender o que determina o Contrato de Concessão, de 19 de setembro de 2008 e seus Anexos III: Regulamento interno dos terminais rodoviários, Anexos V: Programa de manutenção dos terminais rodoviários e Anexo VI: Fiscalização e controle da qualidade dos serviços e outras normas pertinentes como a NR10.",
        )
        adicionar_paragrafo_justificado(
            doc,
            "Informamos também que serão realizadas novas fiscalizações de acompanhamento deste relatório até que todos os itens sejam sanados.",
        )
        adicionar_paragrafo_justificado(
            doc,
            "Segundo a Resolução 83 da Arpe todas as não conformidades devem ser sanadas em um prazo máximo de 180 dias úteis, período em que o processo pode ter arquivamento provisório, sob pena de aplicação de penalidades cabíveis.",
        )

        # =================================================================
        # 5. SEÇÃO DE ASSINATURAS
        # =================================================================
        adicionar_texto_centralizado(
            doc, f'\n\nRecife, {row["Data"]}.'
        )  # Usando a data da sua planilha

        adicionar_texto_centralizado(
            doc, "\n\n\n_______________________________________"
        )
        adicionar_texto_centralizado(doc, "Enildo Manoel da Silva Junior")
        adicionar_texto_centralizado(doc, "Analista de Regulação, matrícula 354-9")

        adicionar_texto_centralizado(doc, "\n\n_______________________________________")
        adicionar_texto_centralizado(doc, "Domingos Sávio Menezes Pereira")
        adicionar_texto_centralizado(
            doc, "Assistente Suplementar de Regulação, matrícula 2581-0"
        )

        adicionar_texto_centralizado(doc, "\n\n_______________________________________")
        adicionar_texto_centralizado(doc, "Maria Ângela Albuquerque de Freitas")
        adicionar_texto_centralizado(
            doc, "Coordenadora de Transportes e Rodovias, matrícula 2590-9"
        )

        # =================================================================
        # 6. SALVAR O DOCUMENTO
        # =================================================================
        nome_relatorio = f"relatorio_{idx_pendente  + 1}.docx"
        caminho_docx = os.path.join(RELATORIOS_DIR, nome_relatorio)
        doc.save(caminho_docx)
        convert(caminho_docx, caminho_docx.replace(".docx", ".pdf"))
        planilha.at[idx_pendente, COLUNA_STATUS] = True
        print(f"✅ Relatório gerado para índice {idx_pendente  + 1}")

    # Antes de salvar, verifica se arquivo está livre
    if arquivo_em_uso(CAMINHO_PLANILHA):
        print(
            f"⚠️ ERRO: O arquivo '{CAMINHO_PLANILHA}' está aberto ou em uso. Não foi possível salvar as alterações."
        )
        return

    # Salva atualização da planilha
    # planilha.to_excel(CAMINHO_PLANILHA, index=False)

    with pd.ExcelWriter(CAMINHO_PLANILHA, engine="openpyxl", mode="w") as writer:
        planilha.to_excel(writer, index=False)

    print("✅ Planilha atualizada com status dos relatórios.")
    ajustar_largura_colunas(CAMINHO_PLANILHA)
    print("🎉 Relatórios pendentes foram gerados com sucesso!")


if __name__ == "__main__":
    try:
        main()
    finally:
        input("\nExecução finalizada. Pressione Enter para sair...")
